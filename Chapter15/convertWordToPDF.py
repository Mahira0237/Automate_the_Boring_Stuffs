# This script runs on Windows only, and you must have Word installed.
import win32com.client 
import docx
wordFilename = 'I:\\My Drive\\AutomateTheBoringStuffs\\Chapter15\\word_document.docx'
pdfFilename = 'I:\\My Drive\\AutomateTheBoringStuffs\\Chapter15\\wordToPDF.pdf'

doc = docx.Document()
doc.add_paragraph('This is a paragraph of text.')
doc.add_heading('Heading 1', 0)
doc.add_heading('Heading 2', 1)
doc.add_heading('Heading 3', 2)
doc.save(wordFilename)

wdFormatPDF = 17 # Word's numeric code for PDFs.
wordObj = win32com.client.Dispatch('Word.Application')

docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()