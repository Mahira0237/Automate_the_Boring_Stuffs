import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
textFile=open('I:\\My Drive\\AutomateTheBoringStuffs\\Chapter15\\CustomInvitationsProject\\guests.txt','r')
names=textFile.readlines()
doc = docx.Document('I:\\My Drive\\AutomateTheBoringStuffs\\Chapter15\\CustomInvitationsProject\\CustomStyle.docx')
for i in range(0,len(names)):
    line1=doc.add_paragraph('It would be a pleasure to have the company of')
    line1.style='custom'
    line1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line2=doc.add_paragraph(names[i].replace('\n',''))
    line2.runs[0].bold=True
    line2.runs[0].font.size=Pt(20)
    line2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line3=doc.add_paragraph('at')
    line3.runs[0].underline=True
    line3.add_run('11010 Memory Lane on the Evening of')
    line3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line3.style='custom'
    line4=doc.add_paragraph('April 1st')
    line4.style='Normal'
    line4.runs[0].font.size=Pt(16)
    line4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line5=doc.add_paragraph('at')
    line5.runs[0].underline=True
    line5.add_run('7o\'clock')
    line5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line5.style='custom'
    doc.add_page_break() 
doc.save('I:\\My Drive\\AutomateTheBoringStuffs\\Chapter15\\CustomInvitationsProject\\invitations.docx')

